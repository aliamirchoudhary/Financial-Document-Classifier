"""Text extraction for the classifier service (self-contained).

Brings over the exact extraction logic from pipeline-2 src/extract.py so the
service has no dependency on the training repo layout:
  - PDF: cheap reading-order first; multiword metric detects garbage (rotated
    scan) -> 4-rotation OCR fallback.
  - tabular (csv/xlsx/xls/xlsm): read as text rows.
  - docx: paragraphs + table cells.
  - images/png/jpg/jfif: routed to the OCR path (liteparse on the image).
Output is a list of {"page_num": int, "text": str}.
"""

from __future__ import annotations

import os
import re
import tempfile
from collections import Counter
from pathlib import Path

import numpy as np
import pandas as pd
import cv2
from PIL import Image
from docx import Document as DocxDocument
import fitz

MIN_CONF = 0.40
GLYPH_REPEAT = 4
READABLE_MULTIWORD = 6
MAX_TABULAR_ROWS = 200
MAX_TABULAR_COLS = 50


_Liteparse = None


def _get_liteparse(max_pages=4, dpi=150):
    global _Liteparse
    if _Liteparse is None:
        from liteparse import LiteParse  # noqa: PLC0415
        _Liteparse = LiteParse(
            ocr_enabled=True,
            ocr_language="eng",
            dpi=dpi,
            quiet=True,
            num_workers=1,
            max_pages=max_pages,
            preserve_very_small_text=True,
        )
    return _Liteparse


def _item_box(item):
    for attr in ["bbox", "bounding_box", "box"]:
        v = getattr(item, attr, None)
        if v is not None:
            return list(v)
    x, y, w, h = (getattr(item, "x", None), getattr(item, "y", None),
                  getattr(item, "width", None), getattr(item, "height", None))
    if all(v is not None for v in (x, y, w, h)):
        return [x, y, x + w, y + h]
    return None


def _multiword_words(items):
    total = 0
    for it in items:
        toks = re.findall(r"[A-Za-z]{3,}", it["text"])
        if len(toks) >= 2:
            total += len(toks)
    return total


def _filter_noise(items):
    if not items:
        return []
    cnt = Counter(it["text"] for it in items)
    keep = []
    for it in items:
        if it["confidence"] < MIN_CONF:
            continue
        if len(it["text"]) <= 2 and cnt[it["text"]] >= GLYPH_REPEAT:
            continue
        keep.append(it)
    return keep


def _render_page(fp, page_num, dpi=150):
    try:
        doc = fitz.open(str(fp))
        page = doc[page_num - 1]
        pix = page.get_pixmap(dpi=dpi)
        doc.close()
    except Exception:
        return None
    img = np.frombuffer(pix.tobytes("png"), dtype=np.uint8)
    img = cv2.imdecode(img, cv2.IMREAD_COLOR)
    if img is None:
        return None
    return cv2.cvtColor(img, cv2.COLOR_BGR2RGB)


def _ocr_image(img):
    fd, tmp = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        Image.fromarray(img).save(tmp)
        res = _get_liteparse(max_pages=1).parse(tmp)
        pg = res.pages[0]
        items = []
        for it in pg.text_items:
            t = getattr(it, "text", "")
            if t and t.strip():
                items.append({
                    "text": t.strip(),
                    "bbox": _item_box(it),
                    "confidence": getattr(it, "confidence", 1.0),
                })
        return items
    finally:
        try:
            os.unlink(tmp)
        except Exception:
            pass


def detect_rotation(fp: Path, sample_pages=None, dpi=150):
    try:
        doc = fitz.open(str(fp))
        npgs = len(doc)
        sample = sample_pages or list(range(1, min(npgs, 2) + 1))
        best, best_conf = 0, 0.0
        for angle in (0, 90, 180, 270):
            tot = 0.0
            for pg in sample:
                img = _render_page(fp, pg, dpi)
                if img is None:
                    continue
                r = img if angle == 0 else np.rot90(img, k=angle // 90)
                items = _filter_noise(_ocr_image(r))
                tot += _multiword_words(items)
            if tot > best_conf:
                best, best_conf = angle, tot
        doc.close()
        return best
    except Exception:
        return 0


def _count_pages(file_path):
    try:
        doc = fitz.open(str(file_path))
        n = len(doc)
        doc.close()
        return n
    except Exception:
        return 1


def _naive_pdf_text(fp, max_pages=4, dpi=150):
    try:
        parse = _get_liteparse(max_pages=max_pages, dpi=dpi)
        result = parse.parse(str(fp))
        pages = []
        for i, page in enumerate(result.pages):
            items = []
            for it in page.text_items:
                t = getattr(it, "text", "")
                if t and t.strip():
                    items.append(t.strip())
            pages.append({"page_num": i + 1, "text": "\n".join(items)})
        return pages
    except Exception:
        return []


def _is_readable(pages):
    strong_lines = 0
    for p in pages:
        for line in p.get("text", "").splitlines():
            toks = re.findall(r"[A-Za-z]{3,}", line)
            if len(toks) >= 2:
                strong_lines += 1
    return strong_lines >= READABLE_MULTIWORD


def _extract_pdf(fp: Path, max_pages=4, dpi=150):
    pages = _naive_pdf_text(fp, max_pages=max_pages, dpi=dpi)
    if _is_readable(pages):
        return pages
    rot = detect_rotation(fp, dpi=dpi)
    if rot == 0:
        return pages
    n = _count_pages(fp)
    out = []
    for pg in range(1, min(n, max_pages) + 1):
        img = _render_page(fp, pg, dpi)
        if img is None:
            out.append({"page_num": pg, "text": ""})
            continue
        r = img if rot == 0 else np.rot90(img, k=rot // 90)
        items = _filter_noise(_ocr_image(r))
        out.append({"page_num": pg, "text": "\n".join(it["text"] for it in items)})
    return out


def _extract_tabular(fp: Path):
    ext = fp.suffix.lower()
    parts = []
    if ext in (".csv", ".tsv"):
        sep = "\t" if ext == ".tsv" else ","
        try:
            df = pd.read_csv(fp, sep=sep, header=None)
        except UnicodeDecodeError:
            df = pd.read_csv(fp, sep=sep, encoding="latin1", header=None)
    else:
        df = pd.DataFrame()
        try:
            df = pd.read_excel(fp, header=None, engine="openpyxl")
            if df.empty:
                xls = pd.ExcelFile(fp, engine="openpyxl")
                for s in xls.sheet_names[1:]:
                    df = pd.read_excel(fp, sheet_name=s, header=None)
                    if not df.empty:
                        break
        except Exception:
            pass
        df = df.head(MAX_TABULAR_ROWS).iloc[:, :MAX_TABULAR_COLS]
    for _, row in df.iterrows():
        vals = [str(v).strip() for v in row if pd.notna(v) and str(v).strip()]
        if vals:
            parts.append(" ".join(vals))
    text = "\n".join(parts)
    return [{"page_num": 1, "text": text}]


def _extract_docx(fp: Path):
    doc = DocxDocument(str(fp))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return [{"page_num": 1, "text": text}]


def extract_text_bytes(filename: str, data: bytes, max_pages=4, dpi=150):
    """Extract page text from raw uploaded bytes (no temp-file name needed).

    Returns list of {"page_num": int, "text": str}. Empty if unreadable.
    """
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        if suffix in (".csv", ".tsv", ".xlsx", ".xls", ".xlsm"):
            return _extract_tabular(tmp_path)
        if suffix == ".docx":
            return _extract_docx(tmp_path)
        return _extract_pdf(tmp_path, max_pages=max_pages, dpi=dpi)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass